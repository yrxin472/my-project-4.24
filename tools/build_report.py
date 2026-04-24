from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import sys
sys.path.append(str(Path(__file__).resolve().parents[1]))

from src.utils import load_json  # noqa: E402


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.bold = True


def add_paragraph(doc: Document, text: str, bold: bool = False, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold


def add_picture_if_exists(doc: Document, path: Path, width: float = 5.8) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph()
        p.add_run(f"[待插入图片] {path.name}").italic = True


def build_report(experiment_dir: Path) -> Path:
    config = load_json(experiment_dir / "config.json")
    stats = load_json(experiment_dir / "stats.json")
    history = load_json(experiment_dir / "history.json")
    train_summary = load_json(experiment_dir / "train_summary.json")
    test_metrics_path = experiment_dir / "test_metrics.json"
    test_metrics = load_json(test_metrics_path) if test_metrics_path.exists() else None

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("从零开始构建三层神经网络分类器实现 EuroSAT 地表覆盖分类\n实验报告")
    run.bold = True
    run.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("（根据作业要求自动生成，可直接补充 GitHub 与权重链接后提交）").italic = True

    add_heading(doc, "1. 作业要求对应情况", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "作业要求"
    hdr[1].text = "本实现对应方式"
    rows = [
        ("自主实现自动微分与反向传播", "src/autograd.py 中实现 Tensor 计算图与 backward()"),
        ("数据加载与预处理", "src/data.py 实现数据扫描、分层划分、均值方差统计、批加载"),
        ("模型定义", "src/model.py 实现三层全连接 MLPClassifier"),
        ("训练循环", "train.py + src/trainer.py 实现 SGD、学习率衰减、最佳权重保存"),
        ("测试评估", "test.py 输出 Accuracy、混淆矩阵、分类别准确率"),
        ("超参数查找", "search.py 支持网格搜索与随机搜索"),
        ("权重可视化", "visualize.py 输出第一层权重图 first_layer_weights.png"),
        ("错例分析", "visualize.py 输出 misclassified_samples.png"),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right

    add_heading(doc, "2. 数据集与预处理", level=1)
    add_paragraph(doc, f"数据集为 EuroSAT_RGB，包含 {len(stats['class_names'])} 个类别。图像统一调整为 {stats['image_size']}×{stats['image_size']} RGB，并展平成向量输入 MLP。")
    add_paragraph(doc, f"训练集通道均值为 {stats['mean']}，通道标准差为 {stats['std']}。")

    add_heading(doc, "3. 模型结构", level=1)
    add_paragraph(
        doc,
        f"模型采用三层全连接结构：FC1 -> 激活 -> FC2 -> 激活 -> FC3。"
        f"其中输入维度为 {config['input_dim']}，隐藏层维度为 {config['hidden_dim']}，"
        f"输出类别数为 {config['num_classes']}，激活函数为 {config['activation']}。"
    )

    add_heading(doc, "4. 训练设置", level=1)
    add_paragraph(
        doc,
        f"优化器为 SGD，初始学习率 {config['lr']}，权重衰减 {config['weight_decay']}，"
        f"学习率衰减策略为每 {config['lr_step']} 个 epoch 衰减为原来的 {config['lr_gamma']} 倍。"
        f"训练轮数 {config['epochs']}，batch size {config['batch_size']}。"
    )
    add_paragraph(doc, f"最佳验证集准确率为 {train_summary['best_val_acc']:.4f}，出现在第 {train_summary['best_epoch']} 个 epoch。")

    add_heading(doc, "5. 训练过程可视化", level=1)
    add_paragraph(doc, "根据作业要求，下面给出训练集/验证集 Loss 曲线以及验证集 Accuracy 曲线。")
    add_picture_if_exists(doc, experiment_dir / "training_curves.png", width=6.0)

    add_heading(doc, "6. 测试集结果", level=1)
    if test_metrics is not None:
        add_paragraph(doc, f"测试集准确率（Accuracy）为 {test_metrics['test_accuracy']:.4f}。")
        add_picture_if_exists(doc, experiment_dir / "confusion_matrix.png", width=5.6)
        add_paragraph(doc, "分类别准确率如下：")
        table2 = doc.add_table(rows=1, cols=2)
        table2.style = "Table Grid"
        table2.rows[0].cells[0].text = "类别"
        table2.rows[0].cells[1].text = "准确率"
        for name, acc in test_metrics["per_class_accuracy"].items():
            row = table2.add_row().cells
            row[0].text = str(name)
            row[1].text = f"{acc:.4f}"
    else:
        add_paragraph(doc, "尚未运行 test.py，因此此处暂缺测试集定量结果。请先执行测试脚本。")

    add_heading(doc, "7. 权重可视化与空间模式观察", level=1)
    add_paragraph(doc, "根据作业要求，将第一层隐藏层权重恢复为图像尺寸进行可视化，并分析其学到的颜色/纹理偏好。")
    add_picture_if_exists(doc, experiment_dir / "first_layer_weights.png", width=6.0)
    add_paragraph(
        doc,
        "分析建议：可重点观察是否存在对绿色纹理、蓝色水体、灰白线性结构等模式敏感的神经元；"
        "若某些权重图表现出明显的条带、边缘或颜色偏置，可将其解释为模型对森林、河流、高速公路等类别早期视觉线索的响应。"
    )

    add_heading(doc, "8. 错例分析（Error Analysis）", level=1)
    add_paragraph(doc, "根据作业要求，选取测试集中若干误分类样本，并结合地物纹理与颜色特征分析原因。")
    add_picture_if_exists(doc, experiment_dir / "misclassified_samples.png", width=6.0)
    add_paragraph(
        doc,
        "常见误分类原因包括：河流与高速公路都可能呈现狭长线性结构；住宅区与工业区都存在规则几何纹理；"
        "草地、农田、森林在局部区域上可能具有相近的绿色纹理分布。"
    )

    add_heading(doc, "9. 超参数搜索说明", level=1)
    add_paragraph(
        doc,
        "本项目支持利用网格搜索或随机搜索对学习率、隐藏层大小、激活函数与 L2 正则强度进行调优。"
        "若已经运行 search.py，请在正文中补充 search_results.csv 的最佳配置与性能变化分析。"
    )

    add_heading(doc, "10. 提交信息（需自行补全）", level=1)
    add_paragraph(doc, "GitHub Repo 链接：<请替换为你的 Public GitHub Repo>")
    add_paragraph(doc, "模型权重下载地址：<请替换为你的 Google Drive / 夸克 / OneDrive 链接>")
    add_paragraph(doc, "README 中应说明环境依赖、训练命令、测试命令与可视化命令。")

    out_path = experiment_dir / "experiment_report.docx"
    doc.save(str(out_path))
    return out_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a Word experiment report from saved training results.")
    parser.add_argument("--experiment_dir", type=str, required=True)
    args = parser.parse_args()
    out_path = build_report(Path(args.experiment_dir))
    print(f"Report saved to: {out_path}")


if __name__ == "__main__":
    main()
